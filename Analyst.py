"""
NOT IN USE ARCHIVED IGNORE

Three-Stage Cerebras AI Financial Data Processor

This script processes Excel financial data through 3 sequential Cerebras AI calls:
1. Identify core important elements in the data
2. Find the specific information we need  
3. Format the data appropriately

Usage: python Analyst.py <excel_file>
"""

import os
import json
import argparse
from typing import Dict, Any, List
import openpyxl
from cerebras.cloud.sdk import Cerebras
from docx import Document  # pip install python-docx
import time
import logging
import random



logger = logging.getLogger(__name__)


class FinancialDataProcessor:
    def __init__(self, api_key: str = None):
        self.client = Cerebras(api_key=api_key or os.environ.get("CEREBRAS_API_KEY"))
        self.model = "qwen-3-coder-480b"
        self._first_api_start = None
        logger.debug("Initialized FinancialDataProcessor with model %s", self.model)
    
    def load_excel_data(self, filepath: str) -> Dict[str, List[List[Any]]]:
        """Load all sheets from Excel workbook into matrices"""
        wb = openpyxl.load_workbook(filepath, data_only=True)
        data = {}
        for ws in wb.worksheets:
            sheet_matrix = []
            for row in ws.iter_rows(values_only=True):
                sheet_matrix.append(list(row))
            data[ws.title] = sheet_matrix
            logger.debug("Loaded sheet '%s' with %d rows", ws.title, len(sheet_matrix))
        return data

    def _chat_with_retries(self, messages: List[Dict[str, Any]], purpose: str, max_retries: int = 3, base_delay: float = 0.5) -> str:
        """Call Cerebras chat with retries and exponential backoff."""
        last_error: Exception | None = None
        for attempt in range(1, max_retries + 1):
            try:
                logger.info("Cerebras call (%s), attempt %d/%d", purpose, attempt, max_retries)
                response = self.client.chat.completions.create(
                    model=self.model,
                    messages=messages
                )
                return response.choices[0].message.content
            except Exception as e:
                last_error = e
                if attempt == 1 and self._first_api_start is None:
                    self._first_api_start = time.perf_counter()
                if attempt == max_retries:
                    break
                sleep_s = base_delay * (2 ** (attempt - 1)) + random.uniform(0, 0.25)
                logger.warning("Cerebras call failed (%s) on attempt %d: %s. Retrying in %.2fs...", purpose, attempt, e, sleep_s)
                time.sleep(sleep_s)
        logger.error("Cerebras call failed after %d attempts (%s)", max_retries, purpose)
        raise last_error if last_error else RuntimeError("Unknown error in _chat_with_retries")
    
    def call_1_identify_core_elements(self, workbook_data: Dict[str, Any]) -> str:
        """
        Cerebras Call 1: Identify core important elements in the financial data
        Returns: Key financial metrics, sheet structure, and data patterns
        """
        system_prompt = """
        You are a financial data analyst. Analyze this Excel workbook and identify the core important elements.
        Provide a clear analysis of the key financial metrics, data patterns, years covered, and data quality.
        Be thorough and analytical in your response.
        """
        
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": json.dumps({"workbook": workbook_data})}
        ]
        
        if self._first_api_start is None:
            self._first_api_start = time.perf_counter()

        return self._chat_with_retries(messages, purpose="identify_core_elements")
    
    def call_2_find_information(self, workbook_data: Dict[str, Any], core_elements: str, question: str = None) -> str:
        """
        Cerebras Call 2: Find the specific information we need based on core elements
        Returns: Extracted values, calculations, and relevant data points
        """
        system_prompt = """
        You are a financial analyst. Based on the core elements identified, extract the specific financial information needed.
        Provide detailed analysis including key financial values, calculations, trends, and actionable recommendations.
        Be precise with numbers and thorough in your analysis.
        """
        
        context = {
            "workbook": workbook_data,
            "core_elements": core_elements,
            "question": question or "Analyze the financial performance and provide key insights"
        }
        
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": json.dumps(context)}
        ]
        
        return self._chat_with_retries(messages, purpose="find_information")
    
    def call_3_format_data(self, workbook_data: Dict[str, Any], core_elements: str, extracted_info: str) -> str:
        """
        Cerebras Call 3: Format the data appropriately for final output
        Returns: Well-structured, formatted financial report
        """
        system_prompt = """
        You are a financial report formatter. Create a professional, well-formatted financial analysis report based on user requests. 
        Be serious and analytic. MAKE SURE YOUR MATH IS CORRECT.
        If the user requests documents such as balance sheets or P&Ls, format them as accurately and professionally as possible.
        Use clear formatting, tables where appropriate.
        """
        
        context = {
            "workbook": workbook_data,
            "core_elements": core_elements,
            "extracted_info": extracted_info
        }
        
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": json.dumps(context)}
        ]
        
        return self._chat_with_retries(messages, purpose="format_data")

    def compute_python_checks(self, workbook_data: Dict[str, List[List[Any]]]) -> str:
        """Compute deterministic numeric summaries per sheet to validate model math."""
        lines: List[str] = []
        for sheet_name, matrix in workbook_data.items():
            numeric_values: List[float] = []
            for row in matrix:
                for cell in row:
                    if isinstance(cell, (int, float)) and cell is not None:
                        numeric_values.append(float(cell))
            count = len(numeric_values)
            total = sum(numeric_values) if count else 0.0
            mean = (total / count) if count else 0.0
            lines.append(f"Sheet: {sheet_name}")
            lines.append(f"  numeric_cells: {count}")
            lines.append(f"  sum: {total:.6f}")
            lines.append(f"  mean: {mean:.6f}")
        if not lines:
            lines.append("No numeric data detected in workbook.")
        return "\n".join(lines)
    
    def process_financial_data(self, filepath: str, question: str = None) -> Dict[str, str]:
        """
        Main orchestration function that runs all 3 Cerebras calls in sequence
        """
        logger.info("Loading Excel data...")
        workbook_data = self.load_excel_data(filepath)
        
        logger.info("Call 1: Identifying core important elements...")
        core_elements = self.call_1_identify_core_elements(workbook_data)
        
        logger.info("Call 2: Finding specific information...")
        extracted_info = self.call_2_find_information(workbook_data, core_elements, question)
        
        logger.info("Call 3: Formatting data...")
        formatted_report = self.call_3_format_data(workbook_data, core_elements, extracted_info)

        # Deterministic Python checks appended to report
        logger.info("Computing deterministic Python checks...")
        python_checks = self.compute_python_checks(workbook_data)
        formatted_report = (
            f"{formatted_report}\n\n=== DETERMINISTIC PYTHON CHECKS ===\n{python_checks}"
        )
        
        return {
            "core_elements": core_elements,
            "extracted_info": extracted_info,
            "formatted_report": formatted_report
        }

def main():
    parser = argparse.ArgumentParser(description="Process financial Excel data with 3-stage Cerebras AI pipeline")
    parser.add_argument("excel_file", help="Path to the Excel workbook to analyze")
    parser.add_argument("--api-key", help="Cerebras API key (overrides env var)")
    parser.add_argument("--output", help="Output JSON file path")
    parser.add_argument("--debug", action="store_true", help="Print detailed debug info")
    parser.add_argument("--log-level", default="INFO", choices=["DEBUG", "INFO", "WARNING", "ERROR", "CRITICAL"], help="Logging level")
    
    args = parser.parse_args()
    
    # Configure logging
    logging.basicConfig(
        level=getattr(logging, args.log_level.upper(), logging.INFO),
        format="%(asctime)s %(levelname)s %(name)s - %(message)s"
    )
    
    # Prompt user for their question
    logger.info("=== Financial Data Processor ===")
    logger.info("Processing: %s", args.excel_file)
    question = input("Enter your financial analysis question: ").strip()
    
    if not question:
        question = "Provide a comprehensive financial analysis of this data"
    
    processor = FinancialDataProcessor(api_key=args.api_key)
    
    try:
        result = processor.process_financial_data(args.excel_file, question)
        
        if args.debug:
            logger.debug("\n=== CORE ELEMENTS ===\n%s", result["core_elements"]) 
            logger.debug("\n=== EXTRACTED INFO ===\n%s", result["extracted_info"]) 
        
        logger.info("\n=== FORMATTED REPORT ===\n%s", result["formatted_report"]) 
        
        if args.output:
            with open(args.output, 'w') as f:
                f.write("=== CORE ELEMENTS ===\n")
                f.write(result["core_elements"])
                f.write("\n\n=== EXTRACTED INFO ===\n")
                f.write(result["extracted_info"])
                f.write("\n\n=== FORMATTED REPORT ===\n")
                f.write(result["formatted_report"])
            logger.info("Full results saved to: %s", args.output)

        # Always write a Word document named 'report.docx' with the formatted output
        try:
            doc = Document()
            doc.add_heading("Financial Analysis Report", level=1)
            for line in result["formatted_report"].split("\n"):
                doc.add_paragraph(line)
            doc.save("report.docx")
            logger.info("Word document written to report.docx")
        except Exception as e:
            logger.warning("Could not write Word document: %s", e)

        # Timing: first API call -> document saved
        if getattr(processor, "_first_api_start", None) is not None:
            elapsed = time.perf_counter() - processor._first_api_start
            logger.info("Time from first API call to document saved: %.3f seconds", elapsed)
            
    except Exception as e:
        logger.exception("Processing failed: %s", e)
        return 1
    
    return 0

if __name__ == "__main__":
    exit(main())
