"""
Three-Stage Cerebras AI Financial Data Processor

This script processes Excel financial data through 3 sequential Cerebras AI calls:
1. Identify core important elements in the data
2. Find the specific information we need  
3. Format the data appropriately

Usage: python new.py <excel_file> [question]
"""

import os
import json
import argparse
from typing import Dict, Any, List
import openpyxl
from cerebras.cloud.sdk import Cerebras
from docx import Document  # pip install python-docx

class FinancialDataProcessor:
    def __init__(self, api_key: str = None):
        self.client = Cerebras(api_key=api_key or os.environ.get("CEREBRAS_API_KEY"))
        self.model = "qwen-3-coder-480b"
    
    def load_excel_data(self, filepath: str) -> Dict[str, List[List[Any]]]:
        """Load all sheets from Excel workbook into matrices"""
        wb = openpyxl.load_workbook(filepath, data_only=True)
        data = {}
        for ws in wb.worksheets:
            sheet_matrix = []
            for row in ws.iter_rows(values_only=True):
                sheet_matrix.append(list(row))
            data[ws.title] = sheet_matrix
        return data
    
    def call_1_identify_core_elements(self, workbook_data: Dict[str, Any]) -> str:
        """
        Cerebras Call 1: Identify core important elements in the financial data
        Returns: Key financial metrics, sheet structure, and data patterns
        """
        system_prompt = """
        You are a financial data analyst. Analyze this Excel workbook and identify the core important elements.
        Provide a clear analysis of the key financial metrics, data patterns, years covered, and data quality.
        Be thorough and analytical in your response.
        """
        
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": json.dumps({"workbook": workbook_data})}
        ]
        
        response = self.client.chat.completions.create(
            model=self.model,
            messages=messages
        )
        
        return response.choices[0].message.content
    
    def call_2_find_information(self, workbook_data: Dict[str, Any], core_elements: str, question: str = None) -> str:
        """
        Cerebras Call 2: Find the specific information we need based on core elements
        Returns: Extracted values, calculations, and relevant data points
        """
        system_prompt = """
        You are a financial analyst. Based on the core elements identified, extract the specific financial information needed.
        Provide detailed analysis including key financial values, calculations, trends, and actionable recommendations.
        Be precise with numbers and thorough in your analysis.
        """
        
        context = {
            "workbook": workbook_data,
            "core_elements": core_elements,
            "question": question or "Analyze the financial performance and provide key insights"
        }
        
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": json.dumps(context)}
        ]
        
        response = self.client.chat.completions.create(
            model=self.model,
            messages=messages
        )
        
        return response.choices[0].message.content
    
    def call_3_format_data(self, workbook_data: Dict[str, Any], core_elements: str, extracted_info: str) -> str:
        """
        Cerebras Call 3: Format the data appropriately for final output
        Returns: Well-structured, formatted financial report
        """
        system_prompt = """
        You are a financial report formatter. Create a professional, well-formatted financial analysis report based on user requests. 
        Be serious and analytic. MAKE SURE YOUR MATH IS CORRECT.
        If the user requests documents such as balance sheets or P&Ls, format them as accurately and professionally as possible.
        Use clear formatting, tables where appropriate.
        """
        
        context = {
            "workbook": workbook_data,
            "core_elements": core_elements,
            "extracted_info": extracted_info
        }
        
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": json.dumps(context)}
        ]
        
        response = self.client.chat.completions.create(
            model=self.model,
            messages=messages
        )
        
        return response.choices[0].message.content
    
    def process_financial_data(self, filepath: str, question: str = None) -> Dict[str, str]:
        """
        Main orchestration function that runs all 3 Cerebras calls in sequence
        """
        print("Loading Excel data...")
        workbook_data = self.load_excel_data(filepath)
        
        print("Call 1: Identifying core important elements...")
        core_elements = self.call_1_identify_core_elements(workbook_data)
        
        print("Call 2: Finding specific information...")
        extracted_info = self.call_2_find_information(workbook_data, core_elements, question)
        
        print("Call 3: Formatting data...")
        formatted_report = self.call_3_format_data(workbook_data, core_elements, extracted_info)
        
        return {
            "core_elements": core_elements,
            "extracted_info": extracted_info,
            "formatted_report": formatted_report
        }

def main():
    parser = argparse.ArgumentParser(description="Process financial Excel data with 3-stage Cerebras AI pipeline")
    parser.add_argument("--api-key", help="Cerebras API key (overrides env var)")
    parser.add_argument("--output", help="Output JSON file path")
    parser.add_argument("--debug", action="store_true", help="Print detailed debug info")
    
    args = parser.parse_args()
    
    # Hardcoded Excel file path
    excel_file = "/Users/ianbaime/Desktop/excel/test.xlsx"
    
    # Prompt user for their question
    print("=== Financial Data Processor ===")
    print(f"Processing: {excel_file}")
    print()
    question = input("Enter your financial analysis question: ").strip()
    
    if not question:
        question = "Provide a comprehensive financial analysis of this data"
    
    processor = FinancialDataProcessor(api_key=args.api_key)
    
    try:
        result = processor.process_financial_data(excel_file, question)
        
        if args.debug:
            print("\n=== CORE ELEMENTS ===")
            print(result["core_elements"])
            print("\n=== EXTRACTED INFO ===")
            print(result["extracted_info"])
        
        print("\n=== FORMATTED REPORT ===")
        print(result["formatted_report"])
        
        if args.output:
            with open(args.output, 'w') as f:
                f.write("=== CORE ELEMENTS ===\n")
                f.write(result["core_elements"])
                f.write("\n\n=== EXTRACTED INFO ===\n")
                f.write(result["extracted_info"])
                f.write("\n\n=== FORMATTED REPORT ===\n")
                f.write(result["formatted_report"])
            print(f"\nFull results saved to: {args.output}")

        # Always write a Word document named 'report.docx' with the formatted output
        try:
            doc = Document()
            doc.add_heading("Financial Analysis Report", level=1)
            for line in result["formatted_report"].split("\n"):
                doc.add_paragraph(line)
            doc.save("report.docx")
            print("Word document written to report.docx")
        except Exception as e:
            print(f"Warning: could not write Word document: {e}")
            
    except Exception as e:
        print(f"Error: {e}")
        return 1
    
    return 0

if __name__ == "__main__":
    exit(main())
